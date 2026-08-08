from __future__ import annotations

import csv
from io import BytesIO
from pathlib import Path
from uuid import UUID, uuid4

import pandas as pd
from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.orm import selectinload
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.rag_chunk import RagChunk
from app.models.rag_document import RagDocument
from app.models.user import User
from app.services.embedding_service import embed_query, embed_texts

SUPPORTED_CONTENT_TYPES = {
    "application/pdf",
    "text/plain",
    "text/markdown",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/csv",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".csv", ".xlsx"}
READY_STATUS = "ready"
FAILED_STATUS = "failed"
PROCESSING_STATUS = "processing"


def _normalize_filename(filename: str) -> str:
    return Path(filename or "document").name


def _ensure_supported_file(filename: str, content_type: str | None) -> None:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS and (content_type or "") not in SUPPORTED_CONTENT_TYPES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type")


def _chunk_text(text: str) -> list[str]:
    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not normalized:
        return []

    chunk_size = max(200, settings.rag_chunk_size)
    overlap = max(0, min(settings.rag_chunk_overlap, chunk_size // 2))
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(len(normalized), start + chunk_size)
        chunks.append(normalized[start:end])
        if end >= len(normalized):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _extract_pdf_text(file_bytes: bytes) -> tuple[str, int | None]:
    reader = PdfReader(BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages), len(reader.pages)


def _extract_docx_text(file_bytes: bytes) -> tuple[str, int | None]:
    document = DocxDocument(BytesIO(file_bytes))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs), None


def _extract_csv_text(file_bytes: bytes) -> tuple[str, int | None]:
    decoded = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(decoded.splitlines())
    rows = [" | ".join(cell.strip() for cell in row if cell.strip()) for row in reader]
    return "\n".join(row for row in rows if row.strip()), len(rows)


def _extract_xlsx_text(file_bytes: bytes) -> tuple[str, int | None]:
    workbook = pd.read_excel(BytesIO(file_bytes), sheet_name=None)
    lines: list[str] = []
    row_count = 0
    for sheet_name, frame in workbook.items():
        lines.append(f"Sheet: {sheet_name}")
        for _, row in frame.fillna("").iterrows():
            values = [str(value).strip() for value in row.tolist() if str(value).strip()]
            if values:
                lines.append(" | ".join(values))
                row_count += 1
    return "\n".join(lines), row_count


def _extract_text(filename: str, content_type: str | None, file_bytes: bytes) -> tuple[str, int | None]:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf" or content_type == "application/pdf":
        return _extract_pdf_text(file_bytes)
    if extension in {".txt", ".md"} or content_type in {"text/plain", "text/markdown"}:
        return file_bytes.decode("utf-8", errors="ignore"), None
    if extension == ".docx":
        return _extract_docx_text(file_bytes)
    if extension == ".csv" or content_type == "text/csv":
        return _extract_csv_text(file_bytes)
    if extension == ".xlsx":
        return _extract_xlsx_text(file_bytes)
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type")


def _build_rag_context(results: list[dict]) -> str:
    lines = [
        "Use the following knowledge base excerpts as primary context when relevant.",
        "If the answer depends on them, cite the document name.",
        "If the knowledge base does not contain the answer, say so clearly.",
    ]
    for index, result in enumerate(results, start=1):
        lines.append(f"{index}. {result['document_name']} (chunk {result['chunk_index']})")
        lines.append(f"   Content: {result['content']}")
    return "\n".join(lines)


async def save_uploaded_document(db: AsyncSession, file: UploadFile, admin_user: User) -> RagDocument:
    filename = _normalize_filename(file.filename or "document")
    _ensure_supported_file(filename, file.content_type)
    file_bytes = await file.read()
    max_size_bytes = settings.rag_max_file_size_mb * 1024 * 1024
    if len(file_bytes) > max_size_bytes:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"File exceeds {settings.rag_max_file_size_mb} MB limit")

    upload_dir = settings.rag_upload_path
    upload_dir.mkdir(parents=True, exist_ok=True)
    stored_name = f"{uuid4().hex}{Path(filename).suffix.lower()}"
    storage_path = upload_dir / stored_name
    storage_path.write_bytes(file_bytes)

    document = RagDocument(
        filename=stored_name,
        original_filename=filename,
        content_type=file.content_type or "application/octet-stream",
        storage_path=str(storage_path),
        size_bytes=len(file_bytes),
        status=PROCESSING_STATUS,
        uploaded_by_user_id=admin_user.id,
    )
    db.add(document)
    await db.flush()

    try:
        text, page_count = _extract_text(filename, file.content_type, file_bytes)
        chunks = _chunk_text(text)
        if not chunks:
            raise ValueError("No readable text found in document")
        embeddings = await embed_texts(chunks)
        for index, (chunk, embedding) in enumerate(zip(chunks, embeddings, strict=False)):
            db.add(
                RagChunk(
                    document_id=document.id,
                    chunk_index=index,
                    content=chunk,
                    token_count=max(1, len(chunk) // 4),
                    metadata_json={"source": filename},
                    embedding=embedding,
                )
            )
        document.page_count = page_count
        document.chunk_count = len(chunks)
        document.status = READY_STATUS
        document.error_message = None
    except Exception as exc:
        document.status = FAILED_STATUS
        document.error_message = str(exc)
    await db.flush()
    return document


async def list_documents(db: AsyncSession) -> list[RagDocument]:
    result = await db.execute(
        select(RagDocument)
        .options(selectinload(RagDocument.uploaded_by))
        .order_by(RagDocument.created_at.desc())
    )
    return list(result.scalars().all())


async def get_document(db: AsyncSession, document_id: UUID) -> RagDocument | None:
    result = await db.execute(
        select(RagDocument)
        .options(selectinload(RagDocument.uploaded_by))
        .where(RagDocument.id == document_id)
    )
    return result.scalar_one_or_none()


async def delete_document(db: AsyncSession, document: RagDocument) -> None:
    await db.execute(delete(RagChunk).where(RagChunk.document_id == document.id))
    await db.delete(document)
    storage_path = Path(document.storage_path)
    if storage_path.exists():
        storage_path.unlink()


async def retrieve_rag_context(db: AsyncSession, query: str, limit: int | None = None) -> tuple[list[dict], str | None]:
    normalized_query = query.strip()
    if not normalized_query:
        return [], None

    query_embedding = await embed_query(normalized_query)
    top_k = max(1, limit or settings.rag_top_k)
    result = await db.execute(
        select(RagChunk, RagDocument)
        .join(RagDocument, RagChunk.document_id == RagDocument.id)
        .where(RagDocument.status == READY_STATUS)
        .order_by(RagChunk.embedding.cosine_distance(query_embedding))
        .limit(top_k)
    )
    rows = result.all()
    citations = [
        {
            "document_id": str(document.id),
            "document_name": document.original_filename,
            "chunk_index": chunk.chunk_index,
            "content": chunk.content,
            "score": None,
        }
        for chunk, document in rows
    ]
    return citations, _build_rag_context(citations) if citations else None